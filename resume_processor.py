import os
import re
import logging
import datetime
import pdfplumber
import PyPDF2
import docx
import docx2txt
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from PIL import Image
import pytesseract
import pypdfium2 as pdfium
import shutil

# Check if tesseract is available in the system PATH
TESSERACT_AVAILABLE = shutil.which("tesseract") is not None

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a file handler to capture errors for debugging
file_handler = logging.FileHandler('resume_processor.log')
file_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
logger.addHandler(file_handler)

# A comprehensive dictionary of common tech skills to filter "noise" from JD
TECH_SKILLS = {
    'python', 'java', 'javascript', 'js', 'react', 'node', 'angular', 'vue', 'html', 'css', 
    'sql', 'nosql', 'mongodb', 'postgresql', 'mysql', 'aws', 'azure', 'gcp', 'docker', 
    'kubernetes', 'git', 'github', 'agile', 'scrum', 'devops', 'machine learning', 'ml', 
    'ai', 'data science', 'pandas', 'numpy', 'tensorflow', 'pytorch', 'c++', 'c#', 'php', 
    'laravel', 'swift', 'kotlin', 'flutter', 'dart', 'typescript', 'ts', 'figma', 'adobe', 
    'ui', 'ux', 'wireframe', 'mockup', 'rest', 'api', 'graphql', 'django', 'flask', 'spring',
    'hibernate', 'excel', 'powerbi', 'tableau', 'spark', 'hadoop', 'linux', 'bash',
    'prototyping', 'interaction design', 'user research', 'usability', 'illustrator', 'photoshop',
    'sketch', 'invision', 'zeplin', 'jira', 'confluence', 'nextjs', 'next.js', 'express',
    'express.js', 'nestjs', 'redis', 'elasticsearch', 'kafka', 'rabbitmq', 'jenkins', 'travis',
    'circleci', 'terraform', 'ansible', 'chef', 'puppet', 'go', 'golang', 'rust', 'ruby',
    'ruby on rails', 'rails', 'scala', 'perl', 'shell', 'powershell', 'objective-c', 'xcode',
    'android studio', 'xamarin', 'ionic', 'react native', 'svelte', 'jquery', 'bootstrap',
    'tailwind', 'tailwindcss', 'sass', 'less', 'webpack', 'babel', 'vite', 'npm', 'yarn',
    'pnpm', 'nuget', 'maven', 'gradle', 'ant', 'sbt', 'sqlite', 'mariadb', 'cassandra',
    'couchbase', 'dynamodb', 'cosmosdb', 'neo4j', 'graphql', 'apollo', 'grpc', 'soap',
    'oauth', 'jwt', 'saml', 'openid', 'auth0', 'firebase', 'supabase', 'amplify', 'heroku',
    'netlify', 'vercel', 'digitalocean', 'linode', 'datadog', 'new relic', 'splunk', 'elk',
    'prometheus', 'grafana', 'kibana', 'logstash', 'sentry', 'jest', 'mocha', 'chai', 'cypress',
    'selenium', 'puppeteer', 'playwright', 'junit', 'pytest', 'rspec', 'cucumber', 'postman',
    'swagger', 'openapi', 'oauth2', 'ci/cd', 'microservices', 'serverless', 'lambda', 'ec2',
    's3', 'rds', 'vpc', 'iam', 'cloudformation', 'ecs', 'eks', 'aks', 'gke', 'cloud run',
    'cloud functions', 'bigquery', 'redshift', 'snowflake', 'databricks', 'airflow', 'luigi',
    'kafka', 'kinesis', 'pubsub', 'sqs', 'sns', 'celery', 'rabbitmq', 'activemq'
}

def download_nltk_data():
    try:
        nltk.data.find('tokenizers/punkt')
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('punkt')
        nltk.download('stopwords')

download_nltk_data()

def ocr_pdf_fallback(file_path):
    """Converts PDF pages to images using pypdfium2 and runs OCR. Robust for image-based PDFs."""
    try:
        if not TESSERACT_AVAILABLE:
            logger.warning("Tesseract binary not found. OCR fallback skipped.")
            return "ERR_TESSERACT_NOT_FOUND"

        logger.info(f"Running OCR fallback on PDF (pypdfium2): {file_path}")
        pdf = pdfium.PdfDocument(file_path)
        full_text = ""
        for i in range(len(pdf)):
            page = pdf[i]
            # Render page to bitmap at 300 DPI
            bitmap = page.render(scale=300/72)
            pil_image = bitmap.to_pil()
            page_text = pytesseract.image_to_string(pil_image)
            if page_text:
                full_text += page_text + "\n"
        pdf.close()
        return full_text.strip()
    except Exception as e:
        logger.error(f"PDF OCR Fallback Error: {str(e)}")
        return ""

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            # Step 1: Try normal text extraction with pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + " "
            except Exception as pdf_e:
                logger.error(f"pdfplumber error: {pdf_e}")

            # Step 2: Try pypdfium2 if text is still too short
            if len(text.strip()) < 50:
                try:
                    # Try PyPDF2 as a standard fallback
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            extracted = page.extract_text()
                            if extracted: text += extracted + " "
                except Exception as pypdf_e:
                    logger.error(f"PyPDF2 error: {pypdf_e}")

            # Step 3: Try pypdfium2 if text is still too short
            if len(text.strip()) < 50:
                try:
                    pdf = pdfium.PdfDocument(file_path)
                    for i in range(len(pdf)):
                        page = pdf[i]
                        text_page = page.get_textpage()
                        extracted = text_page.get_text_range()
                        if extracted:
                            text += extracted + " "
                    pdf.close()
                except Exception as fium_e:
                    logger.error(f"pypdfium2 text error: {fium_e}")
            
            # Step 3: If still little/no text found, attempt OCR fallback
            clean_text = text.replace('\x00', '').strip()
            if len(clean_text) < 30:
                logger.info("PDF has very little extractable text. Attempting OCR fallback...")
                ocr_text = ocr_pdf_fallback(file_path)
                if ocr_text == "ERR_TESSERACT_NOT_FOUND":
                    return "OCR_FAILED_TESSERACT_NOT_FOUND"
                if ocr_text:
                    text = ocr_text
                    
        elif ext == ".docx":
            logger.info(f"Processing DOCX: {file_path}")
            try:
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                # Include tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += "\n" + cell.text
                logger.info(f"python-docx extracted {len(text)} chars")
                if len(text.strip()) < 30:
                    logger.info("python-docx returned little text, trying docx2txt fallback")
                    text = docx2txt.process(file_path)
            except Exception as e:
                logger.error(f"Docx error (python-docx): {e}")
                try: 
                    text = docx2txt.process(file_path)
                    logger.info(f"docx2txt fallback extracted {len(text)} chars")
                except Exception as e2: 
                    logger.error(f"Docx error (docx2txt fallback): {e2}")
                    text = ""
        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            except Exception as e:
                logger.error(f"Text file error: {e}")
                text = ""
        elif ext == ".doc":
            return "UNSUPPORTED_FORMAT_DOC"
        elif ext in [".png", ".jpg", ".jpeg"]:
            if not TESSERACT_AVAILABLE:
                return "OCR_FAILED_TESSERACT_NOT_FOUND"
            try:
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img)
            except Exception as ocr_e:
                logger.error(f"OCR Error: {str(ocr_e)}")
                return "OCR_FAILED_TESSERACT_NOT_FOUND"
        text = text.replace('\x00', '') 
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return ""
    return text.strip()

def preprocess_text(text):
    if not text: return []
    stop_words = set(stopwords.words('english'))
    words = word_tokenize(text.lower())
    return [w for w in words if w.isalnum() and w not in stop_words]

def extract_contact_info(text):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(\+?\d{1,4}[-.\s]?)?(\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}'
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    email = emails[0] if emails else "Not Found"
    phone = "Not Found"
    if phones:
        p = phones[0]
        phone = "".join(p) if isinstance(p, tuple) else p
    return email, phone.strip()

def extract_experience_years(text):
    current_year = datetime.datetime.now().year
    
    # 1. Explicit "X years"
    exp_pattern = r'(\d{1,2}(?:\.\d)?)\+?\s*(?:years?|yrs?|years?\s+of\s+experience)'
    matches = re.findall(exp_pattern, text.lower())
    if matches: 
        try:
            return float(matches[0])
        except ValueError:
            pass
            
    # 2. Calculate from date ranges
    date_pattern = r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,-]*\d{4}|\d{4})\s*(?:-|to|–|—)\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,-]*\d{4}|\d{4}|present|current|now)'
    date_matches = re.findall(date_pattern, text.lower())
    
    if date_matches:
        total_years = 0.0
        
        for start, end in date_matches:
            start_year_match = re.search(r'\d{4}', start)
            if start_year_match:
                start_year = int(start_year_match.group(0))
                
                if end in ['present', 'current', 'now']:
                    end_year = current_year
                else:
                    end_year_match = re.search(r'\d{4}', end)
                    end_year = int(end_year_match.group(0)) if end_year_match else current_year
                    
                diff = end_year - start_year
                if 0 < diff <= 40:
                    total_years += diff
                    
        if total_years > 0:
            return total_years
            
    # 3. Fallback to keywords
    keywords = {
        5.0: ["senior", "lead", "manager", "principal", "expert"], 
        3.0: ["mid-level", "associate", "intermediate"], 
        0.0: ["fresher", "junior", "intern", "trainee", "student"]
    }
    lower_text = text.lower()
    for years, terms in keywords.items():
        for term in terms:
            if term == "senior" and "senior secondary" in lower_text:
                continue
            if re.search(r'\b' + re.escape(term) + r'\b', lower_text):
                return years
    return 0.0

def extract_experience(text):
    current_year = datetime.datetime.now().year
    
    # Explicit "X years"
    exp_pattern = r'(\d{1,2}\+?)\s*(?:years?|yrs?|years?\s+of\s+experience)'
    matches = re.findall(exp_pattern, text.lower())
    if matches: 
        val = matches[0].replace('+', '')
        if val.isdigit() and int(val) >= 5: return "Senior"
        elif val.isdigit() and int(val) >= 2: return "Mid-Level"
        return f"{matches[0]} Years"
        
    # Calculate from date ranges
    date_pattern = r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,-]*\d{4}|\d{4})\s*(?:-|to|–|—)\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,-]*\d{4}|\d{4}|present|current|now)'
    date_matches = re.findall(date_pattern, text.lower())
    
    if date_matches:
        total_years = 0.0
        
        for start, end in date_matches:
            start_year_match = re.search(r'\d{4}', start)
            if start_year_match:
                start_year = int(start_year_match.group(0))
                
                if end in ['present', 'current', 'now']:
                    end_year = current_year
                else:
                    end_year_match = re.search(r'\d{4}', end)
                    end_year = int(end_year_match.group(0)) if end_year_match else current_year
                    
                diff = end_year - start_year
                if 0 < diff <= 40:
                    total_years += diff
                    
        if total_years > 0:
            if total_years >= 5: return "Senior"
            elif total_years >= 2: return "Mid-Level"
            else: return "Entry Level"
            
    # Fallback to keywords with strict word boundaries
    keywords = {
        "Senior": ["senior", "lead", "manager", "principal", "expert"], 
        "Mid-Level": ["mid-level", "associate", "intermediate"], 
        "Fresher": ["fresher", "junior", "intern", "trainee", "student"]
    }
    lower_text = text.lower()
    for level, terms in keywords.items():
        for term in terms:
            # Avoid matching "Senior Secondary"
            if term == "senior" and "senior secondary" in lower_text:
                continue
            if re.search(r'\b' + re.escape(term) + r'\b', lower_text):
                return level
    return "Fresher" # Default to Fresher if nothing else found

def extract_education(text):
    hierarchy = {
        "PhD": 5, "Ph.D": 5, "Doctorate": 5,
        "MBA": 4, "M.Tech": 4, "M.E": 4, "MSc": 4, "M.Sc": 4, "MCA": 4, "Master": 4, "M.Com": 4, "M.A": 4,
        "B.Tech": 3, "B.E": 3, "BSc": 3, "B.Sc": 3, "BCA": 3, "Bachelor": 3, "B.Com": 3, "B.A": 3,
        "Diploma": 2, "High School": 1
    }
    found = []
    lower_text = text.lower()
    for degree in hierarchy.keys():
        # Match with or without dots, case insensitive, word boundaries handled manually
        pattern = re.escape(degree.lower()).replace(r'\.', r'\.?')
        if re.search(r'\b' + pattern + r'(\b|\s|\.)', lower_text):
            found.append(degree)
    
    if not found:
        return "Degree Not Specified"
        
    found.sort(key=lambda d: hierarchy[d], reverse=True)
    highest_degree = found[0]
    
    if hierarchy[highest_degree] == 5: return "PhD"
    elif hierarchy[highest_degree] == 4: return f"Master's ({highest_degree})" if highest_degree != "Master" else "Master's Degree"
    elif hierarchy[highest_degree] == 3: return f"Bachelor's ({highest_degree})" if highest_degree != "Bachelor" else "Bachelor's Degree"
    elif hierarchy[highest_degree] == 2: return "Diploma"
    return highest_degree

def calculate_smart_match(resume_words, job_desc_words, resume_text, jd_text):
    """Calculates a realistic match score focusing on key technical skills."""
    
    # 1. Identify ALL Candidate Skills (using raw text for multi-word skills)
    all_candidate_skills = set()
    lower_resume_text = resume_text.lower()
    for skill in TECH_SKILLS:
        if re.search(r'\b' + re.escape(skill) + r'\b', lower_resume_text):
            all_candidate_skills.add(skill)
            
    # 2. Identify "Real Skills" in JD
    jd_skills = set()
    lower_jd_text = jd_text.lower()
    for skill in TECH_SKILLS:
        if re.search(r'\b' + re.escape(skill) + r'\b', lower_jd_text):
            jd_skills.add(skill)

    # Match Skills
    matched = sorted(list(all_candidate_skills.intersection(jd_skills)))
    missing = sorted(list(jd_skills.difference(all_candidate_skills)))
    
    if not jd_skills: return 0, [], [], sorted(list(all_candidate_skills))
    
    # 3. Skill Match Percentage (Direct accuracy)
    skill_match_percent = (len(matched) / len(jd_skills)) * 100 if jd_skills else 100
    
    # 4. Experience Bonus (For overall score)
    exp_bonus = 0
    jd_exp = extract_experience(jd_text)
    res_exp = extract_experience(resume_text)
    if jd_exp != "Fresher" and jd_exp == res_exp:
        exp_bonus = 10
        
    # 5. Education Bonus
    edu_bonus = 0
    jd_edu = extract_education(jd_text)
    res_edu = extract_education(resume_text)
    if jd_edu != "Degree Not Specified" and res_edu != "Degree Not Specified":
        if any(d in res_edu for d in ["PhD", "Master", "Bachelor"]):
            edu_bonus = 10

    # Total weighted score for ranking
    total_score = round((skill_match_percent * 0.8) + exp_bonus + edu_bonus, 1)
    
    return {
        "total": min(max(total_score, 0), 100),
        "skill_accuracy": round(skill_match_percent, 1),
        "seniority_match": 100 if res_exp == jd_exp else 50 if res_exp != "Fresher" else 20,
        "matched": matched,
        "missing": missing,
        "all_skills": sorted(list(all_candidate_skills))
    }

def generate_interview_questions(matched, missing, exp_input):
    questions = []
    
    # Robustly handle both float years and string levels
    is_senior = False
    is_fresher = False
    
    if isinstance(exp_input, (int, float)):
        if exp_input >= 7: is_senior = True
        elif exp_input < 2: is_fresher = True
    elif isinstance(exp_input, str):
        level = exp_input.lower()
        if "senior" in level or "lead" in level or "manager" in level: is_senior = True
        elif "fresher" in level or "entry" in level or "intern" in level or "0" in level: is_fresher = True

    if is_senior:
        questions.append("Can you describe a time you led a team through a difficult technical transition?")
        questions.append("How do you approach long-term technical debt vs. short-term feature delivery?")
    elif not is_fresher:
        questions.append("Walk me through your most significant professional achievement so far.")
        questions.append("Describe a complex bug you solved recently and the tools you used to debug it.")
    else:
        questions.append("What was your most challenging academic project and how did you approach it?")
        questions.append("Why did you choose your specific field of study/specialization?")

    if missing and missing[0] != "None":
        skill = missing[0].capitalize()
        questions.append(f"We noticed you don't have much experience with {skill}. How would you go about learning it quickly for this role?")
    
    if matched:
        skill = matched[0].capitalize()
        questions.append(f"How have you applied your expertise in {skill} to solve a complex real-world problem?")
    
    return questions

def extract_resume_sections(text):
    sections = {
        "experience": [],
        "education": [],
        "certifications": [],
        "languages": [],
        "skills": [],
        "contact": []
    }
    
    # --- Content-Aware Patterns ---
    email_re = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    phone_re = re.compile(r'(\+?\d{1,4}[\-.\s]?)?(\(?\d{3}\)?[\-.\s]?)?\d{3}[\-.\s]?\d{4}')
    date_range_re = re.compile(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|april|may|june|july|august|september|october|november|december|\d{4})\s*[\-\u2013\u2014to]+\s*(present|current|now|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|april|may|june|july|august|september|october|november|december|\d{4})', re.IGNORECASE)
    degree_re = re.compile(r'\b(ph\.?d|mba|m\.?tech|m\.?e|msc|mca|b\.?tech|b\.?e|bsc|bca|b\.?com|m\.?com|bachelor|master|diploma|high school|b\.?a|m\.?a|engineering|university|college|institute|school of)\b', re.IGNORECASE)
    cert_re = re.compile(r'\b(certified|certification|certificate|aws certified|google certified|microsoft certified|pmp|scrum master|comptia|cisco|oracle|coursera|udemy|edx|nanodegree)\b', re.IGNORECASE)
    lang_re = re.compile(r'\b(english|hindi|tamil|telugu|kannada|malayalam|marathi|bengali|gujarati|punjabi|urdu|french|spanish|german|chinese|mandarin|japanese|korean|arabic|portuguese|italian|russian|dutch|swedish)\b', re.IGNORECASE)
    
    header_mapping = {
        "experience": ["experience", "employment", "work history", "professional experience", "career", "internship", "internships", "projects"],
        "education": ["education", "academic", "qualifications", "academic background", "schooling"],
        "certifications": ["certifications", "certificates", "licenses", "courses", "achievements", "training"],
        "languages": ["languages", "language proficiency"],
        "skills": ["skills", "technical skills", "core competencies", "technologies", "tools", "expertise"]
    }
    
    lines = text.split('\n')
    current_header_section = None
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        clean = stripped.lower()
        
        is_header = False
        if 0 < len(clean.split()) <= 8:
            for key, keywords in header_mapping.items():
                if any(re.search(r'\b' + re.escape(kw) + r'\b', clean) for kw in keywords):
                    current_header_section = key
                    is_header = True
                    break
        
        if is_header:
            continue
        
        classified = False
        
        if email_re.search(stripped) or phone_re.search(stripped):
            if len(stripped) < 80:
                sections["contact"].append(stripped)
                classified = True
        
        if not classified and degree_re.search(stripped):
            sections["education"].append(stripped)
            classified = True
        
        if not classified and cert_re.search(stripped):
            sections["certifications"].append(stripped)
            classified = True
        
        if not classified and lang_re.search(stripped) and len(stripped.split()) <= 8:
            sections["languages"].append(stripped)
            classified = True
        
        if not classified and date_range_re.search(stripped):
            sections["experience"].append(stripped)
            classified = True
        
        if not classified and current_header_section:
            sections[current_header_section].append(stripped)
    
    if not sections["experience"]:
        for line in lines:
            lower_line = line.strip().lower()
            if any(kw in lower_line for kw in ["intern ", "internship", "trainee", "apprentice"]):
                sections["experience"].append(line.strip())
    
    result = {}
    for key in sections:
        items = sections[key]
        seen = set()
        unique = []
        for item in items:
            if item.lower() not in seen:
                seen.add(item.lower())
                unique.append(item)
        if unique:
            result[key] = "\n".join(["\u2022 " + item for item in unique[:15]])
        else:
            result[key] = "Not Specified"
    
    return result

def extract_candidate_name(text, filename):
    # Try filename first
    name = os.path.splitext(filename)[0]
    # Remove UUID pattern if present (8-4-4-4-12 format)
    name = re.sub(r'^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}_', '', name)
    # Remove common resume words and numbers
    name = re.sub(r'resume|cv|profile|\d{4,}|[-_]', ' ', name, flags=re.IGNORECASE)
    clean_name = " ".join(name.split()).title()
    
    if len(clean_name) > 3:
        return clean_name
        
    # Fallback to first line of text
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if len(line) > 2 and len(line) < 30:
            return line.title()
            
    return "Unknown Candidate"

def process_resume(file_path, job_description):
    text = extract_text(file_path)
    filename = os.path.basename(file_path)
    
    if text == "OCR_FAILED_TESSERACT_NOT_FOUND":
        return {
            "filename": filename,
            "name": extract_candidate_name("", filename),
            "sections": {"experience": "Not Specified", "education": "Not Specified", "certifications": "Not Specified", "languages": "Not Specified", "skills": "Not Specified", "contact": "Not Specified"},
            "score": 0, "skills": "None", "all_skills": "None", "missing": "N/A", "experience": "Unknown", "contact": "Not Found", "education": "Not Specified", "summary": "OCR Failed. Tesseract not installed.", "questions": []
        }
        
    if text == "UNSUPPORTED_FORMAT_DOC":
        return {"filename": filename, "score": 0, "skills": "None", "all_skills": "None", "missing": "N/A", "experience": "Unknown", "contact": "N/A", "education": "N/A", "summary": "Unsupported (.doc) format. Please convert to .docx or .pdf.", "questions": [], "sections": {}}
    
    if not text or len(text) < 20:
        return {"filename": filename, "score": 0, "skills": "None", "all_skills": "None", "missing": "N/A", "experience": "Unknown", "contact": "Not Found", "education": "Not Specified", "summary": "Minimal or Unreadable Content", "questions": [], "sections": {"experience": "Not Specified", "education": "Not Specified", "certifications": "Not Specified", "languages": "Not Specified", "skills": "Not Specified", "contact": "Not Specified"}}
        
    resume_words = preprocess_text(text)
    job_desc_words = preprocess_text(job_description)
    
    match_data = calculate_smart_match(resume_words, job_desc_words, text, job_description)
    score = match_data["total"]
    matched = match_data["matched"]
    missing = match_data["missing"]
    all_skills = match_data["all_skills"]
    skill_accuracy = match_data["skill_accuracy"]
    seniority_match = match_data["seniority_match"]
    
    email, phone = extract_contact_info(text)
    exp = extract_experience(text)
    edu = extract_education(text)
    
    questions = generate_interview_questions(matched, missing, exp)
    cand_name = extract_candidate_name(text, filename)
    sections = extract_resume_sections(text)

    return {
        "filename": filename,
        "name": cand_name,
        "sections": sections,
        "score": score,
        "skill_accuracy": skill_accuracy,
        "seniority_match": seniority_match,
        "skills": ", ".join(matched) if matched else "None",
        "all_skills": ", ".join(all_skills) if all_skills else "None",
        "missing": ", ".join(missing) if missing else "None",
        "experience": exp,
        "contact": f"Email: {email}, Phone: {phone}",
        "education": edu,
        "summary": f"Candidate shows {score}% alignment with {exp} experience. Strongest match in {matched[0] if matched else 'core requirements'}.",
        "questions": questions
    }
